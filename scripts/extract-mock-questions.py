# -*- coding: utf-8 -*-
"""从一级教程 docx 提取 450 道自创模拟题 → 平台题库结构"""
import docx, re, json, html
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

SRC = '/mnt/user-data/uploads/GESP_Cpp_Level1.docx'
OUT = '/home/claude/gesp-platform/gesp-platform/data/mock/level1.json'

def iter_blocks(doc):
    for ch in doc.element.body.iterchildren():
        if ch.tag == qn('w:p'): yield ('p', Paragraph(ch, doc))
        elif ch.tag == qn('w:tbl'): yield ('t', Table(ch, doc))

def is_code_table(t):
    if len(t.rows) != 1 or len(t.columns) != 1: return False
    mono = total = 0
    for p in t.cell(0, 0).paragraphs:
        for r in p.runs:
            if (r.text or '').strip():
                total += 1
                if r.font.name and 'consolas' in r.font.name.lower(): mono += 1
    return total > 0 and mono / total > 0.6

def code_text(t):
    lines = [p.text.rstrip() for p in t.cell(0, 0).paragraphs]
    while lines and not lines[0].strip(): lines.pop(0)
    while lines and not lines[-1].strip(): lines.pop()
    pat = re.compile(r'^\s*\d{1,3}\s{2,}')
    if lines and all((not l.strip()) or pat.match(l) for l in lines):
        lines = [pat.sub('', l) for l in lines]
    return '\n'.join(lines)

RE_CH  = re.compile(r'^第\s*(\d+)\s*章\s{2,}')
RE_MC  = re.compile(r'^【模拟题\s*([A-Z]|\d+)[-–](\d+)】\s*(?:[((]([^))]*)[))])?\s*(.*)$')
RE_TF  = re.compile(r'^【判断\s*J[-–](\d+)】\s*(.*)$')
RE_APP = re.compile(r'^【综合\s*Z[-–](\d+)】\s*(?:[((]([^))]*)[))])?\s*(.*)$')
RE_OPT = re.compile(r'^([A-D])\.\s*(.*)$')
RE_REF = re.compile(r'^【参考考点】\s*(.*)$')
RE_EXP = re.compile(r'^【解析】\s*(.*)$')
RE_ANS = re.compile(r'^【答案】\s*(.*)$')

def clean(s): return re.sub(r'\s+', ' ', s).strip()

doc = docx.Document(SRC)
questions = []
cur_ch = 0
cur_set = ''      # 第15章模拟卷套别 A/B/C
q = None          # 进行中的题
field = None      # stem/ref/exp
RE_CODELINE = re.compile(r'^\s*(\d{1,3})\s{3}(.*)$')

def flush():
    global q
    if q: questions.append(q); q = None

for kind, blk in iter_blocks(doc):
    if kind == 'p':
        txt = blk.text.strip()
        if not txt: continue
        if '\t' in blk.text: continue
        m = RE_CH.match(txt)
        if m: flush(); cur_ch = int(m.group(1)); cur_set = ''; continue
        m = RE_MC.match(txt)
        if m:
            flush()
            g1 = m.group(1)
            if g1.isalpha(): cur_set = g1
            q = {'ch': cur_ch, 'set': cur_set if g1.isalpha() else '', 'type': 'mc', 'num': int(m.group(2)), 'difficulty': m.group(3) or '',
                 'stem': clean(m.group(4)), 'code': '', 'options': {}, 'ref': '', 'exp': [], 'answer': ''}
            field = 'stem'; continue
        m = RE_TF.match(txt)
        if m:
            flush()
            q = {'ch': cur_ch, 'set': cur_set if cur_ch == 15 else '', 'type': 'tf', 'num': int(m.group(1)), 'difficulty': '',
                 'stem': clean(m.group(2)), 'code': '', 'options': {}, 'ref': '', 'exp': [], 'answer': ''}
            field = 'stem'; continue
        m = RE_APP.match(txt)
        if m:
            flush()
            q = {'ch': cur_ch, 'set': cur_set if cur_ch == 15 else '', 'type': 'app', 'num': int(m.group(1)), 'difficulty': m.group(2) or '',
                 'stem': clean(m.group(3)), 'code': '', 'options': {}, 'ref': '', 'exp': [], 'answer': ''}
            field = 'stem'; continue
        if q is None: continue
        m = RE_REF.match(txt)
        if m: q['ref'] = clean(re.sub(r'^\*|\*$', '', m.group(1))); field = 'ref'; continue
        m = RE_EXP.match(txt)
        if m:
            field = 'exp'
            if m.group(1).strip(): q['exp'].append(clean(m.group(1)))
            continue
        m = RE_ANS.match(txt)
        if m:
            q['answer'] = clean(m.group(1)); field = None; flush(); continue
        m = RE_OPT.match(txt)
        if m and q['type'] == 'mc' and field in ('stem', 'opt', None):
            q['options'][m.group(1)] = clean(m.group(2)); field = 'opt'; continue
        # 代码行(行号 + 3空格起):还原为 code 字段,保留缩进与换行
        mc_ = RE_CODELINE.match(blk.text.rstrip())
        if mc_ and q is not None:
            line = mc_.group(2)
            if field == 'exp': q['exp'].append('`' + line.strip() + '`')
            else: q['code'] = (q['code'] + '\n' + line) if q['code'] else line
            continue
        # 普通续行
        if field == 'stem': q['stem'] += (' ' + clean(txt))
        elif field == 'exp': q['exp'].append(clean(txt))
        elif field == 'opt':
            # 选项的续行(罕见):并入最后一个选项
            if q['options']:
                k = sorted(q['options'])[-1]; q['options'][k] += ' ' + clean(txt)
        elif field == 'ref': q['ref'] += ' ' + clean(txt)
    else:  # table
        if q is None: continue
        if is_code_table(blk):
            c = code_text(blk)
            if field == 'exp': q['exp'].append('```' + c + '```')
            else: q['code'] = (q['code'] + '\n\n' + c).strip() if q['code'] else c
        else:
            # 普通小表格:转成文字行并入当前字段(综合题里偶有)
            rows = [' | '.join(cell.text.strip() for cell in r.cells) for r in blk.rows]
            t = ' ; '.join(rows)
            if field == 'exp': q['exp'].append(t)
            else: q['stem'] += ' ' + t
flush()

# ---------- 规范化为平台 schema ----------
def norm_answer(q):
    a = q['answer']
    if q['type'] == 'mc':
        m = re.search(r'[A-D]', a)
        return m.group(0) if m else ''
    if q['type'] == 'tf':
        if '√' in a or '对' in a: return '√'
        if '×' in a or '错' in a: return '×'
        return ''
    return a  # app 保留原文

out_q = []
stat = {}
for q in questions:
    a = norm_answer(q)
    exp = '\n'.join(q['exp'])
    if q['ref']: exp = exp + ('\n\n💡 ' + q['ref'] if exp else '💡 ' + q['ref'])
    item = {
        'qid': f"m1-c{q['ch']}{('-' + q.get('set','').lower()) if q.get('set') else ''}-{q['type']}-{q['num']}",
        'chapter': q['ch'], 'type': q['type'], 'num': q['num'],
        'difficulty': q['difficulty'], 'stem': q['stem'], 'code': q['code'],
        'options': q['options'], 'answer': a, 'explanation': exp,
    }
    out_q.append(item)
    stat.setdefault(q['ch'], {'mc': 0, 'tf': 0, 'app': 0})[q['type']] += 1

import os
os.makedirs(os.path.dirname(OUT), exist_ok=True)
json.dump({'level': 1, 'source': '皮爸皮妈讲信奥《GESP C++ 一级备考教程》自创模拟题', 'questions': out_q},
          open(OUT, 'w', encoding='utf-8'), ensure_ascii=False)

total = len(out_q)
print(f'共提取 {total} 题')
for ch in sorted(stat):
    s = stat[ch]; print(f"  第{ch:>2}章: 单选{s['mc']:>2} 判断{s['tf']:>2} 综合{s['app']}")
# 质检
bad_ans = [x['qid'] for x in out_q if x['type'] in ('mc','tf') and not x['answer']]
bad_opt = [x['qid'] for x in out_q if x['type']=='mc' and len(x['options'])!=4]
no_exp  = [x['qid'] for x in out_q if not x['explanation']]
print('缺答案:', bad_ans[:8] or '无')
print('选项≠4:', bad_opt[:8] or '无')
print('缺解析:', no_exp[:8] or '无')
print('带代码题数:', sum(1 for x in out_q if x['code']))
